import xlrd
import xlwt
import docx
document = docx.Document()
i = 1


def read():
    data = xlrd.open_workbook("~/Documents/课程题库复习资料 1905.xls")
    table = data.sheets()[0]
    nrows = table.nrows  # 行数
    # ncols = table.ncols
    for i in range(1, nrows):
        line = table.row_values(i)
        if line[1] == '判断题':
            # line[2] 题干  line[3] 答案 line[7] 正确 line[8] 错误
            # print(line[1])
            judge(line[2], line[3])
        if line[1] == '单选题':
            # line[2] 题干  line[3] 答案    line[7] A line[8] B line[9] C line[10 D]
            single(line[2], line[3], line[7], line[8], line[9], line[10])
            # print(line[1])
        if line[1] == '多选题':
            single(line[2], line[3], line[7], line[8], line[9], line[10])
        #     # line[2] 题干  line[3] 答案 line[7] A line[8] B line[9] C line[10 D]
        #     print(line[1])


def single(question, answer, A, B, C, D):
    global i
    document.add_paragraph(str(i) + "、" + question)
    i += 1
    document.add_paragraph("A."+A)
    document.add_paragraph("B."+B)
    document.add_paragraph("C."+C)
    document.add_paragraph("D." + D)
    document.add_paragraph("答案：" + answer)
    document.add_paragraph()


def judge(question, answer):
    global i
    flag = "正确" if answer == "A" else "错误"
    document.add_paragraph(str(i) + "、" + question)
    i += 1
    # document.add_paragraph('正确   错误')
    document.add_paragraph("答案：" + flag)
    document.add_paragraph()

    # b.add_run('dolor sit amet.')


if __name__ == "__main__":
    read()
    document.save("mayuan/Test.docx")
